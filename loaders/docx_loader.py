# loaders/docx_loader.py

import os
import uuid
import shutil
import subprocess
import tempfile
from docx import Document as DocxDocument
from core import Document, Loader


class DOCXLoader(Loader):
    """
    Loads .docx (native) and .doc (via LibreOffice conversion) files.

    Emits ONE Document per section (heading + content until next heading).
    Tables are emitted as separate Documents with content_type='table'.

    .doc handling:
    - Requires LibreOffice installed (soffice on PATH)
    - Converts to .docx in a temp dir, then reads normally
    - If LibreOffice not found, raises a clear error with install hint
    """

    HEADING_STYLES = {f"Heading {i}" for i in range(1, 10)}

    def __init__(self, min_chars: int = 20):
        self.min_chars = min_chars

    def load(self, source: str) -> list[Document]:
        ext = os.path.splitext(source)[1].lower()

        if ext == ".docx":
            return self._load_docx(source, original_source=source)

        if ext == ".doc":
            converted = self._convert_doc_to_docx(source)
            try:
                return self._load_docx(converted, original_source=source)
            finally:
                # clean up temp conversion
                shutil.rmtree(os.path.dirname(converted), ignore_errors=True)

        raise ValueError(f"Unsupported extension: {ext}. Use .doc or .docx")

    # ------------------------------------------------------------------
    # .doc → .docx conversion via LibreOffice headless
    # ------------------------------------------------------------------
    def _convert_doc_to_docx(self, doc_path: str) -> str:
        soffice = self._find_soffice()
        if soffice is None:
            raise RuntimeError(
                ".doc files require LibreOffice to convert.\n"
                "Install: https://www.libreoffice.org/download/\n"
                "Or manually 'Save As .docx' in Word/LibreOffice and re-run."
            )

        out_dir = tempfile.mkdtemp(prefix="knowledgeos_doc_")
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "docx",
                 "--outdir", out_dir, doc_path],
                check=True, capture_output=True, timeout=60,
            )
        except subprocess.CalledProcessError as e:
            raise RuntimeError(f"LibreOffice conversion failed: {e.stderr.decode(errors='ignore')}")

        base = os.path.splitext(os.path.basename(doc_path))[0]
        converted = os.path.join(out_dir, f"{base}.docx")
        if not os.path.exists(converted):
            raise RuntimeError(f"Conversion produced no output at {converted}")
        return converted

    @staticmethod
    def _find_soffice() -> str | None:
        # PATH first
        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        if soffice:
            return soffice
        # Common Windows install location
        default = r"C:\Program Files\LibreOffice\program\soffice.exe"
        return default if os.path.exists(default) else None

    # ------------------------------------------------------------------
    # Actual .docx parsing
    # ------------------------------------------------------------------
    def _load_docx(self, path: str, original_source: str) -> list[Document]:
        docx = DocxDocument(path)
        docs = []

        current = {"heading": "Introduction", "level": 0, "paragraphs": []}

        def flush():
            if not current["paragraphs"]:
                return
            text = "\n\n".join(current["paragraphs"]).strip()
            if len(text) < self.min_chars:
                return
            docs.append(Document(
                doc_id   = str(uuid.uuid4()),
                content  = text,
                source   = original_source,
                metadata = {
                    "file_type":     "docx",
                    "content_type":  "text",
                    "section_title": current["heading"],
                    "heading_level": current["level"],
                },
            ))

        for para in docx.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style = para.style.name if para.style else ""
            if style in self.HEADING_STYLES:
                flush()
                level = int(style.split()[-1])
                current = {"heading": text, "level": level, "paragraphs": []}
            else:
                current["paragraphs"].append(text)

        flush()

        for table_idx, table in enumerate(docx.tables):
            table_md = self._table_to_markdown(table)
            if table_md:
                docs.append(Document(
                    doc_id   = str(uuid.uuid4()),
                    content  = table_md,
                    source   = original_source,
                    metadata = {
                        "file_type":    "docx",
                        "content_type": "table",
                        "table_index":  table_idx,
                    },
                ))

        return docs

    @staticmethod
    def _table_to_markdown(table) -> str:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append(cells)

        if not rows or len(rows) < 2:
            return ""

        header = rows[0]
        if sum(1 for c in header if c) < 2:
            return ""

        md  = "| " + " | ".join(header) + " |\n"
        md += "| " + " | ".join(["---"] * len(header)) + " |\n"
        for row in rows[1:]:
            while len(row) < len(header):
                row.append("")
            md += "| " + " | ".join(row[:len(header)]) + " |\n"
        return md.strip()